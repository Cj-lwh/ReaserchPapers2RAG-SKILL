"""
Document Text Extraction Module
Extract text from PDF and Word documents for RAG knowledge base

Supports:
- PDF files (.pdf)
- Word documents (.docx, .doc)

Usage:
    python extract_documents.py /path/to/documents
    python extract_documents.py /path/to/documents -o output.json
    python extract_documents.py --zotero  # Auto-detect Zotero storage
"""
import fitz  # pymupdf
from pathlib import Path
from typing import List, Dict, Optional
import json
import argparse
import sys
import os

# Word document support
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed, Word documents will be skipped")


class DocumentExtractor:
    """Document text extractor supporting PDF and Word formats"""

    def __init__(self, doc_dir: str):
        """
        Initialize extractor

        Args:
            doc_dir: Directory containing documents (can be Zotero storage or custom folder)
        """
        self.doc_dir = Path(doc_dir)

    def extract_text_from_pdf(self, pdf_path: Path) -> Optional[Dict]:
        """
        Extract text from a PDF file

        Returns:
            dict with filename, filepath, text, n_pages, metadata, file_type
        """
        try:
            doc = fitz.open(pdf_path)
            text_parts = []

            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")

            metadata = doc.metadata or {}

            result = {
                'filename': pdf_path.name,
                'filepath': str(pdf_path),
                'text': '\n\n'.join(text_parts),
                'n_pages': len(doc),
                'metadata': {
                    'title': metadata.get('title', ''),
                    'author': metadata.get('author', ''),
                    'subject': metadata.get('subject', ''),
                    'keywords': metadata.get('keywords', ''),
                },
                'file_type': 'pdf'
            }

            doc.close()
            return result

        except Exception as e:
            print(f"Error extracting PDF {pdf_path}: {e}")
            return None

    def extract_text_from_word(self, word_path: Path) -> Optional[Dict]:
        """
        Extract text from a Word document

        Returns:
            dict with filename, filepath, text, n_pages (estimated), metadata, file_type
        """
        if not HAS_DOCX:
            print(f"Skipping Word document {word_path.name}: python-docx not installed")
            return None

        try:
            doc = Document(str(word_path))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            # Try to extract document properties
            try:
                core_props = doc.core_properties
                metadata = {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                }
            except:
                metadata = {}

            total_text = '\n\n'.join(text_parts)
            estimated_pages = max(1, len(total_text) // 3000)

            result = {
                'filename': word_path.name,
                'filepath': str(word_path),
                'text': total_text,
                'n_pages': estimated_pages,
                'metadata': metadata,
                'file_type': 'word'
            }

            return result

        except Exception as e:
            print(f"Error extracting Word document {word_path}: {e}")
            return None

    def extract_text_from_file(self, file_path: Path) -> Optional[Dict]:
        """Extract text based on file type"""
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            return self.extract_text_from_word(file_path)
        else:
            return None

    def extract_all_documents(self) -> List[Dict]:
        """Extract all supported documents from the directory"""
        pdf_files = list(self.doc_dir.rglob('*.pdf'))
        word_files = []
        if HAS_DOCX:
            word_files = list(self.doc_dir.rglob('*.docx')) + list(self.doc_dir.rglob('*.doc'))

        all_files = pdf_files + word_files
        print(f"Found {len(pdf_files)} PDF files, {len(word_files)} Word files")

        results = []
        for i, file_path in enumerate(all_files):
            print(f"Processing [{i+1}/{len(all_files)}]: {file_path.name}")
            result = self.extract_text_from_file(file_path)
            if result and result['text'].strip():
                results.append(result)

        print(f"Successfully extracted {len(results)} documents")
        return results

    def save_to_json(self, results: List[Dict], output_path: str):
        """Save extracted results to JSON file"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)

        print(f"Saved to {output_path}")


def find_zotero_storage() -> Optional[Path]:
    """
    Try to auto-detect Zotero storage path

    Returns:
        Zotero storage path, or None if not found
    """
    # Common Windows paths
    windows_paths = [
        Path.home() / 'Zotero',
        Path.home() / 'Documents' / 'Zotero',
        Path(os.environ.get('APPDATA', '')) / 'Zotero' / 'Zotero',
    ]

    for base_path in windows_paths:
        storage_path = base_path / 'storage'
        if storage_path.exists():
            return storage_path

    return None


def main():
    """Main function: extract text from all documents"""
    parser = argparse.ArgumentParser(
        description='Extract text from PDF and Word documents for RAG knowledge base'
    )
    parser.add_argument(
        'doc_path',
        nargs='?',
        help='Document directory path. Leave empty to auto-detect Zotero storage'
    )
    parser.add_argument(
        '-o', '--output',
        help='Output JSON file path (default: ./extracted_texts.json)'
    )
    parser.add_argument(
        '--zotero',
        action='store_true',
        help='Force use of Zotero storage path'
    )

    args = parser.parse_args()

    # Determine document directory
    if args.doc_path:
        doc_dir = Path(args.doc_path)
        print(f"Using specified path: {doc_dir}")
    elif args.zotero:
        doc_dir = find_zotero_storage()
        if doc_dir is None:
            print("Error: Could not auto-detect Zotero storage path")
            print("Please specify path manually:")
            print("  python extract_documents.py /path/to/your/documents")
            sys.exit(1)
        print(f"Using Zotero storage path: {doc_dir}")
    else:
        doc_dir = find_zotero_storage()
        if doc_dir:
            print(f"Auto-detected Zotero storage path: {doc_dir}")
        else:
            print("Zotero storage path not found. Please specify document directory:")
            print("Usage: python extract_documents.py /path/to/your/documents")
            sys.exit(1)

    if not doc_dir.exists():
        print(f"Error: Directory does not exist: {doc_dir}")
        sys.exit(1)

    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = Path.cwd() / 'extracted_texts.json'

    print(f"Document directory: {doc_dir}")
    print(f"Output file: {output_path}")

    # Extract documents
    extractor = DocumentExtractor(doc_dir)
    results = extractor.extract_all_documents()

    if len(results) == 0:
        print("Warning: No document content extracted")
        return

    extractor.save_to_json(results, output_path)

    # Print statistics
    total_chars = sum(len(r['text']) for r in results)
    pdf_count = sum(1 for r in results if r['file_type'] == 'pdf')
    word_count = sum(1 for r in results if r['file_type'] == 'word')

    print(f"\nStatistics:")
    print(f"  PDF documents: {pdf_count}")
    print(f"  Word documents: {word_count}")
    print(f"  Total documents: {len(results)}")
    print(f"  Total characters: {total_chars:,}")
    print(f"  Average characters: {total_chars // len(results):,}")


if __name__ == '__main__':
    main()
